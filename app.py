import streamlit as st
from docx import Document
import os
from datetime import datetime

st.set_page_config(page_title="DOCX File Merger", layout="centered")

# Function to merge multiple DOCX files
def combine_docs(uploaded_files):
    output_doc = Document()
    for file in uploaded_files:
        doc = Document(file)
        for para in doc.paragraphs:
            output_doc.add_paragraph(para.text)
        output_doc.add_page_break()  # Optional: add page break between files
    # Save file
    filename = f"combined_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_doc.save(filename)
    return filename

# UI starts here
st.title("📄 DOCX File Merger")
st.markdown("Upload multiple `.docx` files and merge them into a single document.")

# Upload section
uploaded_files = st.file_uploader("Upload your .docx files", type=["docx"], accept_multiple_files=True)

# Display selected files
if uploaded_files:
    st.subheader("Selected Files:")
    for file in uploaded_files:
        st.write(f"✅ {file.name}")

    if st.button("🚀 Generate Combined File"):
        with st.spinner("Processing..."):
            output_path = combine_docs(uploaded_files)
            with open(output_path, "rb") as f:
                st.success("✅ File generated successfully!")
                st.download_button("⬇️ Download Combined DOCX", f, file_name=output_path)
            os.remove(output_path)  # Clean up after download
